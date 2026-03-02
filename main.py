from typing import List, Optional, Dict, Any
import re
from pathlib import Path

from pydantic import BaseModel, Field
from dotenv import load_dotenv

load_dotenv()
from langchain_core.messages import HumanMessage, ToolMessage, AIMessage
from langchain_core.runnables import RunnableLambda, RunnableBranch
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain_tavily import TavilySearch
from docx import Document


class JobSearchResult(BaseModel):
    """Schema for a single job posting found in search"""
    title: str = Field(description="Job title")
    company: str = Field(description="Company name")
    location: str = Field(description="Job location")
    url: str = Field(description="Full LinkedIn job URL (e.g., https://linkedin.com/jobs/view/...)")
    description: Optional[str] = Field(default=None, description="Brief job description if available")


class JobSearchResponse(BaseModel):
    """Structured output for job search - returns job URLs directly, no scraping needed"""
    jobs: List[JobSearchResult] = Field(
        description="List of job postings found. Each must include the complete LinkedIn URL."
    )
    search_summary: str = Field(
        description="Brief summary of the search performed and results found"
    )


class JobDescription(BaseModel):
    """Schema for extracted job description"""
    title: str = Field(description="Job title")
    company: str = Field(description="Company name")
    location: str = Field(description="Job location")
    description: str = Field(description="Full job description")
    requirements: str = Field(description="Key requirements and qualifications")
    responsibilities: str = Field(description="Key responsibilities")


class ProfileSkillsUpdate(BaseModel):
    """Schema for profile and skills updates"""
    profile_lines: List[str] = Field(description="Two lines to add to profile section")
    skills: List[str] = Field(description="Skills to add that match the job description")


class JobDescriptionWithUpdates(BaseModel):
    """Combined schema for job description extraction and resume updates in a single call"""
    # Job description fields
    title: str = Field(description="Job title")
    company: str = Field(description="Company name")
    location: str = Field(description="Job location")
    description: str = Field(description="Full job description")
    requirements: str = Field(description="Key requirements and qualifications")
    responsibilities: str = Field(description="Key responsibilities")
    # Resume update fields
    profile_lines: List[str] = Field(description="Exactly 2 lines to add to profile section that align with the job")
    skills: List[str] = Field(description="Skills to add that match the job description")


def find_tool_by_name(tools: List[Any], tool_name: str) -> Any:
    """Find a tool by its name"""
    for tool in tools:
        if hasattr(tool, 'name') and tool.name == tool_name:
            return tool
    raise ValueError(f"Tool with name {tool_name} not found")




def create_job_description_chain(llm: ChatOpenAI, search_tool: TavilySearch):
    """Create LCEL chain for fetching and extracting job description"""
    
    # Step 1: Generate search queries from URL
    def generate_search_queries(inputs: Dict[str, Any]) -> Dict[str, Any]:
        linkedin_url = inputs["linkedin_url"]
        print(f"\n🔍 Fetching job description from: {linkedin_url}")
        job_id = linkedin_url.split("/")[-1] if "/" in linkedin_url else ""
        search_queries = [
            f"LinkedIn job posting {linkedin_url} full description",
            f"site:linkedin.com/jobs {job_id}",
            f"{linkedin_url} job description requirements",
        ]
        return {**inputs, "search_queries": search_queries, "job_id": job_id}
    
    # Step 2: Search and aggregate content
    def search_and_aggregate(inputs: Dict[str, Any]) -> Dict[str, Any]:
        search_queries = inputs["search_queries"]
        page_content = ""
        
        for query in search_queries:
            try:
                result = search_tool.invoke({"query": query, "search_depth": "advanced"})
                
                # Handle different result formats from TavilySearch
                if isinstance(result, list):
                    for item in result:
                        if isinstance(item, dict):
                            content = item.get("content", "") or item.get("raw_content", "")
                            if content:
                                page_content += f"\n\n{content}"
                        else:
                            page_content += f"\n\n{str(item)}"
                elif isinstance(result, dict):
                    content = result.get("content", "") or result.get("raw_content", "")
                    if content:
                        page_content += f"\n\n{content}"
                else:
                    page_content += f"\n\n{str(result)}"
                
                # If we got substantial content, break
                if len(page_content) > 500:
                    break
            except Exception as e:
                print(f"⚠️  Search query failed: {query} - {str(e)}")
                continue
        
        if not page_content or len(page_content.strip()) < 100:
            print("⚠️  Warning: Limited content extracted. The job description may be incomplete.")
        
        return {**inputs, "page_content": page_content[:10000]}
    
    # Step 3: Create extraction prompt
    extraction_prompt_template = ChatPromptTemplate.from_messages([
        ("human", """Extract the job description information from the following LinkedIn job posting content.

URL: {linkedin_url}

Content:
{page_content}

Please extract:
1. Job title
2. Company name
3. Location
4. Full job description
5. Key requirements and qualifications (skills, experience, education)
6. Key responsibilities

If any information is missing, indicate "Not specified".""")
    ])
    
    # Step 4: Extract structured output
    llm_structured = llm.with_structured_output(JobDescription)
    
    # Compose the chain
    chain = (
        RunnableLambda(generate_search_queries)
        | RunnableLambda(search_and_aggregate)
        | extraction_prompt_template
        | llm_structured
    )
    
    return chain


def read_docx_resume(resume_path: str) -> Document:
    """Read DOCX resume file"""
    path = Path(resume_path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {resume_path}")
    
    if path.suffix.lower() not in ['.docx', '.doc']:
        raise ValueError(f"Unsupported file format: {path.suffix}. Only .docx files are supported.")
    
    return Document(resume_path)


def find_section_paragraphs(doc: Document, section_keywords: List[str]) -> Optional[int]:
    """Find the index of the first paragraph in a section"""
    for i, para in enumerate(doc.paragraphs):
        para_text_lower = para.text.lower().strip()
        for keyword in section_keywords:
            if keyword.lower() in para_text_lower:
                return i
    return None


def add_profile_lines(doc: Document, profile_lines: List[str]):
    """Add 2 lines to the beginning of the profile section"""
    print("📝 Adding profile lines...")
    
    # Find profile section (look for common keywords)
    profile_keywords = ['Profile', 'Summary', 'Professional Summary', 'Career Summary']
    profile_idx = find_section_paragraphs(doc, profile_keywords)
    
    if profile_idx is None:
        print("⚠️  Profile section not found. Adding at the beginning.")
        profile_idx = 0
    
    # Find the first content paragraph after the section header
    # Skip empty paragraphs and look for the first paragraph with actual content
    insert_idx = profile_idx + 1
    section_keywords = ['experience', 'skills', 'education', 'work history', 'employment', 'projects']
    
    # Find the first non-empty paragraph that's not another section header
    first_content_idx = None
    for i in range(insert_idx, len(doc.paragraphs)):
        para_text = doc.paragraphs[i].text.strip()
        para_text_lower = para_text.lower()
        
        # Check if this is a new section header
        if any(keyword in para_text_lower for keyword in section_keywords) and len(para_text) < 50:
            # We've reached the next section, stop here
            break
        
        # Found first content paragraph
        if para_text and len(para_text) > 10:  # Has meaningful content
            first_content_idx = i
            break
    
    # If we found a content paragraph, insert before it
    # Otherwise, insert after the header
    if first_content_idx is not None:
        insert_idx = first_content_idx
    else:
        # No content found, insert right after header
        insert_idx = profile_idx + 1
    
    # Get formatting from the target paragraph or a nearby paragraph
    source_para = None
    if insert_idx < len(doc.paragraphs):
        source_para = doc.paragraphs[insert_idx]
    elif profile_idx < len(doc.paragraphs):
        source_para = doc.paragraphs[profile_idx]
    
    # Insert the new lines at the beginning of the Profile section as bullet points
    # Insert in reverse order so they appear in the correct order
    for line in reversed(profile_lines):
        new_para = doc.paragraphs[insert_idx].insert_paragraph_before(line)
        
        # Format as bullet point
        new_para.style = 'List Bullet' if 'List Bullet' in [s.name for s in doc.styles] else None
        # If List Bullet style doesn't exist, manually add bullet character
        if new_para.style is None:
            new_para.text = f"• {line}"
        
        # Copy formatting from source paragraph if available
        if source_para and source_para.runs:
            source_run = source_para.runs[0]
            for run in new_para.runs:
                run.font.name = source_run.font.name
                if source_run.font.size:
                    run.font.size = source_run.font.size
                # Copy other formatting attributes
                run.font.bold = source_run.font.bold
                run.font.italic = source_run.font.italic
    
    print(f"✅ Added {len(profile_lines)} lines to the beginning of profile section")


def add_skills(doc: Document, new_skills: List[str]):
    """Add matching skills to the skills section"""
    print("📝 Adding skills...")
    
    # Find skills section
    skills_keywords = ['Skills']
    skills_idx = find_section_paragraphs(doc, skills_keywords)
    
    if skills_idx is None:
        print("⚠️  Skills section not found. Adding at the end.")
        skills_idx = len(doc.paragraphs) - 1
    
    # Find where to insert (after the section header)
    insert_idx = skills_idx + 1
    
    # Get existing skills text to avoid duplicates
    existing_skills_text = ""
    for i in range(insert_idx, min(insert_idx + 10, len(doc.paragraphs))):
        existing_skills_text += doc.paragraphs[i].text.lower() + " "
    
    # Filter out skills that already exist
    skills_to_add = []
    for skill in new_skills:
        if skill.lower() not in existing_skills_text:
            skills_to_add.append(skill)
    
    if not skills_to_add:
        print("ℹ️  All skills already exist in resume.")
        return
    
    # Find the last paragraph in skills section or insert after header
    last_skills_para = insert_idx
    section_keywords = ['experience', 'education', 'work history', 'employment', 'projects']
    for i in range(insert_idx, len(doc.paragraphs)):
        para_text_lower = doc.paragraphs[i].text.lower().strip()
        if any(keyword in para_text_lower for keyword in section_keywords) and len(para_text_lower) < 50:
            last_skills_para = i
            break
        last_skills_para = i + 1
    
    # Add skills as bullet points
    # Check if existing skills are in a list format
    if insert_idx < len(doc.paragraphs):
        existing_format = doc.paragraphs[insert_idx].text
        if existing_format.strip():
            # Add to existing paragraph if it's a skills list
            if len(existing_format) < 200:  # Likely a skills paragraph
                # Add as comma-separated if existing format is comma-separated
                if "," in existing_format:
                    doc.paragraphs[insert_idx].text += f", {', '.join(skills_to_add)}"
                else:
                    # Add as bullet points
                    for skill in reversed(skills_to_add):
                        new_para = doc.paragraphs[last_skills_para].insert_paragraph_before(f"• {skill}")
                        new_para.style = 'List Bullet' if 'List Bullet' in [s.name for s in doc.styles] else None
                        if doc.paragraphs[insert_idx].runs:
                            source_run = doc.paragraphs[insert_idx].runs[0]
                            for run in new_para.runs:
                                run.font.name = source_run.font.name
                                if source_run.font.size:
                                    run.font.size = source_run.font.size
            else:
                # Create new paragraphs as bullet points
                for skill in reversed(skills_to_add):
                    new_para = doc.paragraphs[last_skills_para].insert_paragraph_before(f"• {skill}")
                    new_para.style = 'List Bullet' if 'List Bullet' in [s.name for s in doc.styles] else None
                    if doc.paragraphs[insert_idx].runs:
                        source_run = doc.paragraphs[insert_idx].runs[0]
                        for run in new_para.runs:
                            run.font.name = source_run.font.name
                            if source_run.font.size:
                                run.font.size = source_run.font.size
        else:
            # Empty paragraph, add as bullet points
            for skill in reversed(skills_to_add):
                new_para = doc.paragraphs[insert_idx].insert_paragraph_before(f"• {skill}")
                new_para.style = 'List Bullet' if 'List Bullet' in [s.name for s in doc.styles] else None
    else:
        # Add at end as bullet points
        for skill in reversed(skills_to_add):
            new_para = doc.add_paragraph(f"• {skill}")
            new_para.style = 'List Bullet' if 'List Bullet' in [s.name for s in doc.styles] else None
    
    print(f"✅ Added {len(skills_to_add)} new skills: {', '.join(skills_to_add)}")


def create_profile_skills_chain(llm: ChatOpenAI):
    """Create LCEL chain for generating profile lines and skills"""
    
    print("🤖 Generating profile lines and skills...")
    
    prompt_template = ChatPromptTemplate.from_messages([
        ("human", """You are a professional resume writer. Based on the following job description, generate:
1. Exactly 2 lines to add to the PROFILE/SUMMARY section that align with the job requirements
2. A list of skills from the job description that should be added to the SKILLS section

JOB DESCRIPTION:
Title: {title}
Company: {company}
Location: {location}

Requirements:
{requirements}

Responsibilities:
{responsibilities}

Full Description:
{description}

CURRENT RESUME (for context):
{existing_resume_text}

INSTRUCTIONS:
- Generate exactly 2 concise, professional lines for the profile section that highlight relevant experience/qualifications matching the job
- Extract skills from the job description that are technical/important (e.g., LangChain, Python, RAG, LLMs, etc.)
- Only include skills that are explicitly mentioned or strongly implied in the job description
- Do NOT include skills that are already in the resume
- Keep profile lines professional and ATS-friendly
- Profile lines should be 1-2 sentences each, maximum 150 characters per line""")
    ])
    
    # Format inputs for the prompt
    def format_inputs(inputs: Dict[str, Any]) -> Dict[str, Any]:
        job_desc = inputs["job_desc"]
        existing_resume_text = inputs["existing_resume_text"]
        return {
            "title": job_desc.title,
            "company": job_desc.company,
            "location": job_desc.location,
            "requirements": job_desc.requirements,
            "responsibilities": job_desc.responsibilities,
            "description": job_desc.description[:2000],
            "existing_resume_text": existing_resume_text[:2000],
        }
    
    llm_structured = llm.with_structured_output(ProfileSkillsUpdate)
    
    chain = (
        RunnableLambda(format_inputs)
        | prompt_template
        | llm_structured
    )
    
    return chain


def create_combined_job_and_updates_chain(llm: ChatOpenAI, search_tool: TavilySearch):
    """
    Create optimized LCEL chain that combines job description extraction 
    and profile/skills generation in a SINGLE LLM call.
    
    This reduces LLM calls from 2 per job to 1 per job (50% reduction).
    """
    
    # Step 1: Generate search queries from URL
    def generate_search_queries(inputs: Dict[str, Any]) -> Dict[str, Any]:
        linkedin_url = inputs["linkedin_url"]
        print(f"\n🔍 Fetching job description from: {linkedin_url}")
        job_id = linkedin_url.split("/")[-1] if "/" in linkedin_url else ""
        search_queries = [
            f"LinkedIn job posting {linkedin_url} full description",
            f"site:linkedin.com/jobs {job_id}",
            f"{linkedin_url} job description requirements",
        ]
        return {**inputs, "search_queries": search_queries, "job_id": job_id}
    
    # Step 2: Search and aggregate content
    def search_and_aggregate(inputs: Dict[str, Any]) -> Dict[str, Any]:
        search_queries = inputs["search_queries"]
        page_content = ""
        
        for query in search_queries:
            try:
                result = search_tool.invoke({"query": query, "search_depth": "advanced"})
                
                # Handle different result formats from TavilySearch
                if isinstance(result, list):
                    for item in result:
                        if isinstance(item, dict):
                            content = item.get("content", "") or item.get("raw_content", "")
                            if content:
                                page_content += f"\n\n{content}"
                        else:
                            page_content += f"\n\n{str(item)}"
                elif isinstance(result, dict):
                    content = result.get("content", "") or result.get("raw_content", "")
                    if content:
                        page_content += f"\n\n{content}"
                else:
                    page_content += f"\n\n{str(result)}"
                
                # If we got substantial content, break
                if len(page_content) > 500:
                    break
            except Exception as e:
                print(f"⚠️  Search query failed: {query} - {str(e)}")
                continue
        
        if not page_content or len(page_content.strip()) < 100:
            print("⚠️  Warning: Limited content extracted. The job description may be incomplete.")
        
        return {**inputs, "page_content": page_content[:10000]}
    
    # Step 3: Create combined extraction and generation prompt
    combined_prompt_template = ChatPromptTemplate.from_messages([
        ("human", """You are a professional resume writer. Extract job description information AND generate resume updates in one pass.

URL: {linkedin_url}

Job Posting Content:
{page_content}

Current Resume (for context):
{existing_resume_text}

TASK 1 - Extract the following job description information:
1. Job title
2. Company name
3. Location
4. Full job description
5. Key requirements and qualifications (skills, experience, education)
6. Key responsibilities

TASK 2 - Generate resume updates based on the job description:
1. Exactly 2 lines to add to the PROFILE/SUMMARY section that align with the job requirements
2. A list of skills from the job description that should be added to the SKILLS section

INSTRUCTIONS FOR RESUME UPDATES:
- Generate exactly 2 concise, professional lines for the profile section that highlight relevant experience/qualifications matching the job
- Extract skills from the job description that are technical/important (e.g., LangChain, Python, RAG, LLMs, etc.)
- Only include skills that are explicitly mentioned or strongly implied in the job description
- Do NOT include skills that are already in the resume
- Keep profile lines professional and ATS-friendly
- Profile lines should be 1-2 sentences each, maximum 150 characters per line

If any job information is missing, indicate "Not specified".""")
    ])
    
    # Step 4: Single LLM call for both extraction and generation
    llm_structured = llm.with_structured_output(JobDescriptionWithUpdates)
    
    # Compose the optimized chain
    chain = (
        RunnableLambda(generate_search_queries)
        | RunnableLambda(search_and_aggregate)
        | combined_prompt_template
        | llm_structured
    )
    
    return chain


def create_resume_update_chain(llm: ChatOpenAI, search_tool: TavilySearch):
    """
    Create optimized LCEL chain for complete resume update workflow.
    
    OPTIMIZED LLM Call Pattern (per job):
    - 1 call: Extract job description AND generate profile/skills in single call (structured output)
    Total: 1 LLM call per job (50% reduction from previous 2 calls)
    
    This combines job description extraction and resume update generation into
    a single LLM call, significantly reducing API costs and latency.
    """
    
    # Create optimized combined chain
    combined_chain = create_combined_job_and_updates_chain(llm, search_tool)
    
    # Step 1: Read resume first (needed for context in the combined call)
    def read_resume(inputs: Dict[str, Any]) -> Dict[str, Any]:
        resume_path = inputs["resume_path"]
        print(f"\n{'='*70}")
        print(f"Processing job: {inputs['job_url']}")
        print(f"{'='*70}")
        print(f"\n📄 Reading resume from: {resume_path}")
        doc = read_docx_resume(resume_path)
        existing_resume_text = "\n".join([para.text for para in doc.paragraphs])
        return {**inputs, "doc": doc, "existing_resume_text": existing_resume_text}
    
    # Step 2: Combined extraction and generation in single LLM call
    def extract_and_generate(inputs: Dict[str, Any]) -> Dict[str, Any]:
        print("🤖 Extracting job description and generating resume updates (single LLM call)...")
        result = combined_chain.invoke({
            "linkedin_url": inputs["job_url"],
            "existing_resume_text": inputs["existing_resume_text"],
        })
        
        print(f"\n✅ Extracted Job Description:")
        print(f"   Title: {result.title}")
        print(f"   Company: {result.company}")
        print(f"   Location: {result.location}")
        print(f"✅ Generated {len(result.profile_lines)} profile lines and {len(result.skills)} skills")
        
        # Split the combined result into job_desc-like object and updates
        # We'll use the result directly since it has both fields
        return {**inputs, "combined_result": result}
    
    # Step 3: Update resume document
    def update_document(inputs: Dict[str, Any]) -> Dict[str, Any]:
        doc = inputs["doc"]
        result = inputs["combined_result"]
        add_profile_lines(doc, result.profile_lines)
        add_skills(doc, result.skills)
        return inputs
    
    # Step 4: Save resume
    def save_resume(inputs: Dict[str, Any]) -> str:
        doc = inputs["doc"]
        result = inputs["combined_result"]
        resume_path = inputs["resume_path"]
        input_path = Path(resume_path)
        company_name = result.company.lower().replace(" ", "_").replace(",", "")
        output_path = input_path.parent / f"{input_path.stem}_aligned_{company_name}{input_path.suffix}"
        doc.save(str(output_path))
        print(f"\n✅ Updated resume saved to: {output_path}")
        return str(output_path)
    
    # Compose the optimized chain
    chain = (
        RunnableLambda(read_resume)
        | RunnableLambda(extract_and_generate)
        | RunnableLambda(update_document)
        | RunnableLambda(save_resume)
    )
    
    return chain


def search_jobs(query: str, llm: ChatOpenAI, search_tool: TavilySearch, max_iterations: int = 5) -> JobSearchResponse:
    """
    Job search agent using LCEL patterns with structured output.
    
    This function:
    - Uses structured output directly (no URL scraping from text)
    - Returns JobSearchResponse with URLs in the schema
    - Uses clean LCEL patterns for agent steps
    - Separated from resume processing pipeline
    """
    
    llm_with_tools = llm.bind_tools([search_tool])
    llm_structured = llm.with_structured_output(JobSearchResponse)
    
    # Create search prompt
    search_prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a job search assistant. Search for job postings and return them with complete LinkedIn URLs.

IMPORTANT:
- Use the search tool to find job postings
- Return results using structured output with complete LinkedIn URLs
- Each job must include: title, company, location, and full LinkedIn URL
- Prefer recent postings when available, but include any relevant matches
- Be thorough but efficient with tool calls"""),
        ("human", "{query}")
    ])
    
    messages = [HumanMessage(content=search_prompt.format(query=query))]
    total_tool_calls = 0
    
    # Agent loop using LCEL patterns
    for iteration in range(max_iterations):
        print(f"\n--- Search Iteration {iteration + 1} ---")
        
        # Call LLM with tools
        ai_message = llm_with_tools.invoke(messages)
        tool_calls = getattr(ai_message, "tool_calls", None) or []
        
        messages.append(ai_message)
        
        if tool_calls:
            # Execute tools
            for tool_call in tool_calls:
                tool_name = tool_call.get("name")
                tool_args = tool_call.get("args", {})
                tool_call_id = tool_call.get("id")
                
                observation = search_tool.invoke(tool_args)
                total_tool_calls += 1
                print(f"🔧 Tool {tool_name} executed (total: {total_tool_calls})")
                
                messages.append(
                    ToolMessage(content=str(observation), tool_call_id=tool_call_id)
                )
        else:
            # No more tool calls - get structured output
            result = llm_structured.invoke(messages)
            print(f"\n✅ Job search complete!")
            print(f"   Found {len(result.jobs)} job(s)")
            print(f"   Total tool calls: {total_tool_calls}")
            return result
    
    # Max iterations reached - still try to get structured output
    print(f"\n⚠️  Max iterations reached. Getting final result...")
    result = llm_structured.invoke(messages)
    return result


def main():
    """
    Main function with proper separation of concerns:
    1. Job Search Agent (uses structured output, no URL scraping)
    2. Resume Processing Pipeline (pure LCEL chain)
    """
    print("🚀 Resume Aligner - Job Search & Resume Processing")
    print("=" * 70)
    
    llm = ChatOpenAI(model="gpt-4o", temperature=0)
    search_tool = TavilySearch()
    
    # ============================================================
    # STEP 1: Job Search Agent (separated from pipeline)
    # ============================================================
    print("\n📋 Step 1: Searching for jobs...")
    print("-" * 70)
    
    # Search query - removed unenforceable "2 days ago" requirement
    query = "Find 3 job postings for an Angular lead developer in Bangalore on LinkedIn"
    
    try:
        search_result = search_jobs(query, llm, search_tool, max_iterations=5)
    except Exception as e:
        print(f"❌ Job search failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return
    
    if not search_result.jobs:
        print("⚠️  No jobs found. Exiting.")
        return
    
    # Extract URLs directly from structured output (no scraping needed)
    job_urls = [job.url for job in search_result.jobs if job.url]
    
    if not job_urls:
        print("⚠️  No valid job URLs found in search results. Exiting.")
        return
    
    print(f"\n✅ Found {len(job_urls)} job(s) with valid URLs:")
    for i, job in enumerate(search_result.jobs, 1):
        print(f"   {i}. {job.title} at {job.company} - {job.url}")
    
    # ============================================================
    # STEP 2: Resume Processing Pipeline (pure LCEL chain)
    # ============================================================
    print(f"\n{'='*70}")
    print("Step 2: Processing jobs and updating resume...")
    print(f"{'='*70}")
    
    resume_path = "Anand-Hiremath_12+Exp_Resume.docx"
    if not Path(resume_path).exists():
        print(f"❌ Resume file not found: {resume_path}")
        print("   Please ensure the resume file exists in the current directory.")
        return
    
    # Create resume processing pipeline (separate from agent)
    resume_update_chain = create_resume_update_chain(llm, search_tool)
    updated_resumes = []
    
    for i, job_url in enumerate(job_urls, 1):
        try:
            print(f"\n📝 Processing job {i}/{len(job_urls)}...")
            output_path = resume_update_chain.invoke({
                "job_url": job_url,
                "resume_path": resume_path,
            })
            updated_resumes.append(output_path)
        except Exception as e:
            print(f"❌ Error processing job {i}: {str(e)}")
            import traceback
            traceback.print_exc()
            continue
    
    # ============================================================
    # Summary
    # ============================================================
    print(f"\n{'='*70}")
    print("✅ Resume Update Complete!")
    print(f"{'='*70}")
    print(f"Updated {len(updated_resumes)} resume(s):")
    for path in updated_resumes:
        print(f"   - {path}")


if __name__ == "__main__":
    main()